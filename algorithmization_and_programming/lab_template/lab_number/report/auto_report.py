import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

# Create a new Document
doc = Document()

# Declare directories
current_dir = os.path.dirname(os.path.abspath(__file__))
project_dir = os.path.dirname(current_dir)

# Construct the full path to files
data_txt_path = os.path.join(current_dir, 'data.txt')
condition_txt_path = os.path.join(current_dir, 'task_condition/data.txt')
unit_result_txt_path = os.path.join(current_dir, 'tests_result/unit_result.txt')

main_cpp_path = os.path.join(project_dir, 'src/main.cpp')
functions_cpp_path = os.path.join(project_dir, 'lib/function.cpp')
unit_tests_cpp_path = os.path.join(project_dir, 'tests/unit_tests.cpp')


print("Start creating report file...")

######################
## Header of report ##
######################

print("Adding header of report...")

# Construct the full path to data.txt
data_txt_path = os.path.join(current_dir, 'data.txt')

# Read the file with data about lab 
with open(data_txt_path, 'r', encoding='utf-8') as file:
    # Read each line and assign to variables
    lab_number = file.readline().strip()
    lab_theme_name = file.readline().strip()
    lab_goal = file.readline().strip()

report_header_text = (
        'ЗВІТ\n'
        f'про виконання лабораторної роботи № {lab_number}\n'
        f'{lab_theme_name}\n'
        'з дисципліни\n'
        '"Алгоритмізація та програмування"\n'
        'студента групи РІ-12\n'
        'Грушевського Івана Олександровича\n'
        )

report_header = doc.add_paragraph()
report_header_run = report_header.add_run(report_header_text)
report_header.alignment = WD_ALIGN_PARAGRAPH.CENTER

report_header_run.font.name = 'Times New Roman'
report_header_run.font.size = Pt(14)
report_header_run.bold = True

doc.add_page_break()


####################
## Task Condition ##
#################### 

print("Adding task condition...")

doc.add_heading(f'Мета: {lab_goal}')
doc.add_heading('Умова завдання:', level=1)

# Read the file with data about lab
with open(condition_txt_path, 'r', encoding='utf-8') as file:
    # Read the entire file content
    condition_text = file.read()

task_condition_paragraph = doc.add_paragraph()

task_condition_run = task_condition_paragraph.add_run(condition_text)
task_condition_run.font.name = 'Times New Roman'
task_condition_run.font.size = Pt(12)

# Add an image of task condition 
task_condition_img_path = os.path.join(current_dir, 'task_condition/image.png')

doc.add_picture(task_condition_img_path, width=Inches(2))

doc.add_page_break()


####################### 
## Structure Diagram ##
####################### 

print("Adding structure diagram...")

doc.add_heading('Структурна схема програми:', level=1)

structure_img_path = condition_img_path = os.path.join(current_dir, 'structure_uml/image.png')

doc.add_picture(structure_img_path, width=Inches(2))

doc.add_page_break()


#################### 
## Programme Text ##
#################### 

print("Adding programme text...")

doc.add_heading('Текст програми:', level=1)

print("main.cpp ...")
doc.add_heading('main.cpp', level=2)

# Read main.cpp file
with open(main_cpp_path, 'r', encoding='utf-8') as file:
    # Read the entire script 
    main_text = file.read()

main_cpp = doc.add_paragraph()

main_cpp_run = main_cpp.add_run(main_text)
main_cpp_run.font.name = 'Courier New'
main_cpp_run.font.size = Pt(10)


doc.add_page_break()

print("function.cpp ...")
doc.add_heading('function.cpp', level=2)

# Read function.cpp file
with open(functions_cpp_path, 'r', encoding='utf-8') as file:
    # Read the entire script 
    function_text = file.read()

function_cpp = doc.add_paragraph()

function_cpp_run = function_cpp.add_run(function_text)
function_cpp_run.font.name = 'Courier New'
function_cpp_run.font.size = Pt(10)

doc.add_page_break()


##############
## Git Repo ##
############## 

print("Adding Git repo link...")

doc.add_heading('Посилання на git-репозиторій з проктом', level=2)

git_link = 'https://github.com/p-i-r-u-m/University-labs/tree/master/AP'

git_link_paragraph = doc.add_paragraph() 

git_link_run = git_link_paragraph.add_run(git_link)
git_link_run.font.name = 'Courier New'
git_link_run.font.size = Pt(10)

doc.add_page_break()


################ 
## Unit Tests ##
################

print("Adding unit tests...")

doc.add_heading('Результати unit-тесту:')

doc.add_heading('unit_tests.cpp', level=2)

# Read unit_tests.cpp file
with open(unit_tests_cpp_path, 'r', encoding='utf-8') as file:
    # Read the entire script 
    unit_text = file.read()

unit_tests_cpp = doc.add_paragraph()

unit_tests_cpp_run = unit_tests_cpp.add_run(unit_text)
unit_tests_cpp_run.font.name = 'Courier New'
unit_tests_cpp_run.font.size = Pt(10)


doc.add_page_break()


doc.add_heading('Вивід unit-тесту:', level=2)

# Read unit_result.txt file
with open(unit_result_txt_path, 'r', encoding='utf-8') as file:
    # Read the entire script 
    result_text = file.read()

unit_result = doc.add_paragraph()

unit_result_run = unit_result.add_run(result_text)
unit_result_run.font.name = 'Courier New'
unit_result_run.font.size = Pt(10)

doc.add_page_break()


#################
## Conclusions ##
#################

print("Adding conclusion...")

doc.add_heading('Висновки:')

conclusion_text = f'У результаті виконання лабораторної роботи я зміг {lab_goal}'

conclusion_paragraph = doc.add_paragraph()

conclusion_run = conclusion_paragraph.add_run(conclusion_text)
conclusion_run.font.name = 'Times New Roman'
conclusion_run.font.size = Pt(12)


# Save the Document
doc.save(f'{current_dir}/АП_РІ-12_Грушевський_ЛР-{lab_number}.docx')

print(f"\nCreated АП_РІ-12_Грушевський_ЛР-{lab_number}.docx")
